import io
import os
import fitz  # PyMuPDF
import docx

def extract_resume_text(file_bytes: bytes, filename: str) -> str:
    """
    Extract text from resume PDF or DOCX file bytes.
    """
    _, ext = os.path.splitext(filename)
    ext = ext.lower()
    
    if ext == ".pdf":
        text = ""
        # Open PDF from bytes
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        for page in doc:
            page_text = page.get_text()
            if page_text:
                text += page_text + "\n"
        return text.strip()
        
    elif ext == ".docx":
        doc = docx.Document(io.BytesIO(file_bytes))
        text_parts = []
        
        # Extract paragraph text
        for paragraph in doc.paragraphs:
            if paragraph.text:
                text_parts.append(paragraph.text)
                
        # Extract text from tables to ensure multi-column templates are parsed
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text and cell_text not in text_parts:
                        text_parts.append(cell_text)
                        
        return "\n".join(text_parts).strip()
        
    else:
        raise ValueError(f"Unsupported file extension: {ext}")
